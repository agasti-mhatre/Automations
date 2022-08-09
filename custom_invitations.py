import docx 


#find all of the names in the text file
text_file = open("guests.txt", "r")
names = []
for line in text_file.readlines():
    names.append(line.replace('\n', ''))

text_file.close()


#create word document and create invitations
page_number = 0
doc = docx.Document()
for name in range(0, len(names)):
    doc.add_paragraph("It would be a pleasure to have the company of")
    doc.paragraphs[(page_number * 5)].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[(page_number * 5)].style = "Intense Quote"

    #insert name of the person
    doc.add_paragraph("{}".format(names[name]))
    doc.paragraphs[(page_number * 5) + 1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[(page_number * 5) + 1].style = "Quote"

    doc.add_paragraph("at")
    doc.paragraphs[(page_number * 5) + 2].add_run(" 11010 Memory Lane on the Evening of")
    doc.paragraphs[(page_number * 5) + 2].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    

    doc.add_paragraph("April 1st")
    doc.paragraphs[(page_number * 5) + 3].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


    doc.add_paragraph("at")
    doc.paragraphs[(page_number * 5) + 4].add_run(" 7 o'clock")
    doc.paragraphs[(page_number * 5) + 4].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    
    #don't break after the last page
    if name >= len(names) - 1:
        break

    doc.paragraphs[(page_number * 5) + 4].runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)
    page_number += 1



doc.save("invitations.docx")